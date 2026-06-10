"""Tests for the document-HTML sanitizer (critical XSS guard).

These pin the property that untrusted markdown/docx content cannot smuggle a
script, event handler, or javascript:/data: navigation into the WebView DOM, while
the formatting and the app's own structural markup survive. If the sanitize call
were removed from textdoc/officedoc, the markdown/docx/export tests here fail.
"""

import os
import tempfile

import pytest

from renderers import textdoc
from renderers.sanitize import sanitize_html

# ── unit: sanitize_html allowlist behavior ──────────────────────────────────

ATTACKS = [
    "<img src=x onerror=alert(1)>",
    "<svg onload=alert(2)></svg>",
    "<script>alert(3)</script>",
    "<iframe src=http://evil></iframe>",
    "<object data=evil></object>",
    "<a href=\"javascript:alert(4)\">x</a>",
    "<a href=\"vbscript:msgbox(5)\">x</a>",
    "<p style=\"background:url(javascript:alert(6))\">x</p>",
    "<a href=\"data:text/html;base64,PHNjcmlwdD5hbGVydCg3KTwvc2NyaXB0Pg==\">x</a>",
    "<img src=\"data:text/html,<script>alert(8)</script>\">",
    "<form action=\"javascript:alert(9)\"><button>go</button></form>",
    "<body onload=alert(10)>",
    "<details ontoggle=alert(11) open>x</details>",
    "<a href=\"  JaVaScRiPt:alert(12)\">x</a>",   # scheme obfuscation
]


@pytest.mark.parametrize("payload", ATTACKS)
def test_sanitize_strips_attacks(payload):
    out = sanitize_html(payload).lower()
    assert "onerror" not in out
    assert "onload" not in out
    assert "ontoggle" not in out
    assert "<script" not in out
    assert "<svg" not in out
    assert "<iframe" not in out
    assert "<object" not in out
    assert "javascript:" not in out
    assert "vbscript:" not in out
    assert "data:text/html" not in out
    assert "style=" not in out


def test_sanitize_keeps_data_image_only_on_img():
    # docx/pptx inline pictures as base64 data:image — must survive on <img src>.
    img = '<img src="data:image/png;base64,iVBORw0KGgo=" alt="pic">'
    assert "data:image/png" in sanitize_html(img)
    # but a navigable data: link must be dropped even if it is an image type.
    link = '<a href="data:image/png;base64,iVBORw0KGgo=">x</a>'
    assert "data:image" not in sanitize_html(link)


def test_sanitize_keeps_structural_markup():
    html = (
        '<h2 id="sec">S</h2>'
        '<sup class="fn-ref" data-fn-id="1" contenteditable="false">1</sup>'
        '<section class="doc-footnotes" data-doc-footnotes="1">'
        '<ol class="fn-list"><li class="fn-item" data-fn-id="1">note</li></ol></section>'
        '<table><tbody><tr><td colspan="2">c</td></tr></tbody></table>'
        '<a href="https://ok.example/">link</a>'
        '<a href="#sec">anchor</a>'
    )
    out = sanitize_html(html)
    assert 'id="sec"' in out                 # TOC anchor target
    assert 'data-fn-id="1"' in out           # footnote linkage
    assert 'data-doc-footnotes="1"' in out
    assert 'contenteditable="false"' in out
    assert 'colspan="2"' in out
    assert 'https://ok.example/' in out
    assert 'href="#sec"' in out              # internal nav still works


def test_sanitize_strips_data_tx_orig():
    # The frontend translation tool restores data-tx-orig via innerHTML, so a
    # document must NOT be able to supply that attribute (XSS bypass otherwise).
    out = sanitize_html('<p data-tx-orig="<img src=x onerror=alert(1)>">hi</p>')
    assert "data-tx-orig" not in out
    assert "onerror" not in out


def test_sanitize_is_idempotent():
    once = sanitize_html("".join(ATTACKS))
    assert sanitize_html(once) == once


def test_sanitize_empty_input():
    assert sanitize_html("") == ""
    assert sanitize_html(None) == ""


def test_sanitize_fail_closed_when_nh3_absent(monkeypatch):
    # If nh3 is missing in a frozen build, we must escape (not pass raw HTML).
    import renderers.sanitize as s
    monkeypatch.setattr(s, "_HAVE_NH3", False)
    out = s.sanitize_html("<img src=x onerror=alert(1)>")
    # Everything is escaped to inert text: no live tags, angle brackets encoded.
    assert "<img" not in out
    assert "<" not in out and ">" not in out
    assert "&lt;img" in out


# ── integration: the actual render paths are sanitized ──────────────────────

_MD_ATTACK = (
    "# Heading\n\n"
    "<img src=x onerror=alert('md')>\n\n"
    "<script>alert('s')</script>\n\n"
    "[evil](javascript:alert('a'))\n\n"
    "[safe](https://example.com)\n\n"
    "## Second\n\n| a | b |\n|---|---|\n| 1 | 2 |\n"
)


def test_markdown_render_is_sanitized():
    r = textdoc.render_text(_MD_ATTACK, "evil.md", ".md")
    h = r["html"].lower()
    assert "onerror" not in h
    assert "<script" not in h
    assert "javascript:" not in h
    # formatting + navigation survive
    assert 'id="heading"' in h
    assert "https://example.com" in h
    assert "<table>" in h
    assert [o["anchor"] for o in r["outline"]] == ["heading", "second"]


def test_export_html_is_sanitized():
    out = textdoc.build_export_html(_MD_ATTACK, "evil.md")
    low = out.lower()
    assert "onerror" not in low
    assert "<script" not in low
    assert "javascript:" not in low
    assert "<!doctype html>" in low          # wrapper intact
    assert "https://example.com" in out


def test_docx_javascript_hyperlink_is_stripped():
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    tmp = tempfile.mkdtemp(prefix="yr_xss_")
    try:
        path = os.path.join(tmp, "evil.docx")
        d = docx.Document()
        d.add_heading("Doc Title", level=1)
        p = d.add_paragraph("Click ")
        r_id = p.part.relate_to(
            "javascript:alert(document.title)",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyper = p._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
        run = p.add_run("the link")
        hyper.append(run._r)
        p._p.append(hyper)
        d.add_paragraph("Body text survives.")
        d.save(path)

        from renderers import officedoc
        res = officedoc._docx_to_html(path)
        h = res["html"]
        assert "javascript:" not in h
        assert "Doc Title" in h           # content preserved
        assert "Body text survives." in h
        assert 'id="h-1"' in h            # heading anchor preserved
    finally:
        import shutil
        shutil.rmtree(tmp, ignore_errors=True)
